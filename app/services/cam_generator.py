from docx import Document

def generate_cam(company, financials, ratios, risk):
    doc = Document()
    doc.add_heading("Credit Appraisal Memo", level=1)
    doc.add_paragraph(f"Company: {company}")
    doc.add_heading("Financial Summary", level=2)
    for k,v in financials.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("Financial Ratios", level=2)
    for k,v in ratios.items():
        doc.add_paragraph(f"{k}: {round(v,3)}")
    doc.add_heading("Risk Assessment", level=2)
    doc.add_paragraph(f"Risk Score: {risk['score']}")
    doc.add_paragraph(f"Risk Level: {risk['risk_level']}")
    file_path = f"uploads/{company}_CAM.docx"
    doc.save(file_path)
    return file_path
