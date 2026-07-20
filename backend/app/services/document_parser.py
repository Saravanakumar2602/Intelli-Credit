import pdfplumber
import docx
import pandas as pd
import os

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"[PARSER ERROR] PDF extraction failed for {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables inside the Word document
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        text = "\n".join(full_text)
    except Exception as e:
        print(f"[PARSER ERROR] DOCX extraction failed for {file_path}: {e}")
    return text

def extract_text_from_xlsx(file_path):
    text = ""
    try:
        xl = pd.ExcelFile(file_path)
        full_text = []
        for sheet_name in xl.sheet_names:
            full_text.append(f"--- Sheet: {sheet_name} ---")
            df = xl.parse(sheet_name)
            # Use to_string to maintain grid representation
            full_text.append(df.to_string(index=False))
        text = "\n\n".join(full_text)
    except Exception as e:
        print(f"[PARSER ERROR] XLSX extraction failed for {file_path}: {e}")
    return text

def extract_text_from_file(file_path):
    """Unified text extraction routing based on file extension"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in {".docx", ".doc"}:
        return extract_text_from_docx(file_path)
    elif ext in {".xlsx", ".xls"}:
        return extract_text_from_xlsx(file_path)
    else:
        print(f"[PARSER WARNING] Unsupported extension {ext} for {file_path}. Returning empty text.")
        return ""

