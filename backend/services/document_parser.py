import os
import pdfplumber
import docx
import pandas as pd
import pytesseract
from PIL import Image
from backend.core.config import settings

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF. If the PDF is scanned (extracts no text),
    falls back to pytesseract OCR by rendering pages to images natively.
    """
    text = ""
    try:
        # Set tesseract command path if specified
        if settings.TESSERACT_CMD != "tesseract":
            pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
            
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text and len(page_text.strip()) > 50:
                    text += page_text + "\n"
                else:
                    # Scanned PDF page fallback - OCR page image using pytesseract
                    print(f"[PARSER OCR] Page {i+1} of {file_path} appears scanned. Triggering Tesseract OCR...")
                    try:
                        # Render page to image using pdfplumber native rendering
                        img_obj = page.to_image(resolution=150)
                        pil_img = img_obj.original # PIL Image object
                        ocr_text = pytesseract.image_to_string(pil_img)
                        if ocr_text:
                            text += ocr_text + "\n"
                    except Exception as ocr_err:
                        print(f"[PARSER OCR ERROR] Page {i+1} OCR failed: {ocr_err}")
    except Exception as e:
        print(f"[PARSER ERROR] PDF extraction failed for {file_path}: {e}")
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text and tables from Word documents"""
    text = ""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        text = "\n".join(full_text)
    except Exception as e:
        print(f"[PARSER ERROR] DOCX extraction failed for {file_path}: {e}")
    return text

def extract_text_from_xlsx(file_path: str) -> str:
    """Extract tables and spreadsheets retaining grid representation"""
    text = ""
    try:
        xl = pd.ExcelFile(file_path)
        full_text = []
        for sheet_name in xl.sheet_names:
            full_text.append(f"--- Sheet: {sheet_name} ---")
            df = xl.parse(sheet_name)
            # Use to_string to maintain grid representation
            full_text.append(df.to_string(index=False))
        text = "\n\n".join(full_text)
    except Exception as e:
        print(f"[PARSER ERROR] XLSX extraction failed for {file_path}: {e}")
    return text

def extract_text_from_csv(file_path: str) -> str:
    """Extract and format csv spreadsheet grids"""
    try:
        df = pd.read_csv(file_path)
        return df.to_string(index=False)
    except Exception as e:
        print(f"[PARSER ERROR] CSV extraction failed for {file_path}: {e}")
        return ""

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files"""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"[PARSER ERROR] TXT extraction failed for {file_path}: {e}")
        return ""

def extract_text_from_file(file_path: str) -> str:
    """Unified file text extraction router routing by extension"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in {".docx", ".doc"}:
        return extract_text_from_docx(file_path)
    elif ext in {".xlsx", ".xls"}:
        return extract_text_from_xlsx(file_path)
    elif ext == ".csv":
        return extract_text_from_csv(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    else:
        print(f"[PARSER WARNING] Unsupported extension {ext} for {file_path}. Returning empty text.")
        return ""
